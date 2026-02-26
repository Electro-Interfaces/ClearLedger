"""Локальный парсинг файлов (Layer 1a) — только механический, без AI."""

from __future__ import annotations

import csv
import io
from typing import Any


def parse_file(content: bytes, mime_type: str, file_name: str) -> dict[str, Any]:
    """Извлекает текст и метаданные из файла.

    Returns:
        {"extracted_text": str, "extracted_fields": dict, "page_count": int | None}
    """
    if mime_type == "application/pdf":
        return _parse_pdf(content)
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _parse_excel(content)
    if mime_type in ("text/xml", "application/xml"):
        return _parse_xml(content)
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(content)
    if mime_type == "text/csv":
        return _parse_csv(content)
    if mime_type.startswith("text/"):
        return _parse_text(content)
    # Скан/фото — просто сохраняем, текст не извлекаем
    return {"extracted_text": "", "extracted_fields": {}, "page_count": None}


def _parse_pdf(content: bytes) -> dict[str, Any]:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=content, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    text = "\n\n".join(pages)
    return {
        "extracted_text": text,
        "extracted_fields": {},
        "page_count": len(doc),
    }


def _parse_excel(content: bytes) -> dict[str, Any]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheets = {}
    row_count = 0
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(c) if c is not None else "" for c in row])
            row_count += 1
        sheets[name] = rows[:5]  # первые 5 строк для preview
    wb.close()

    return {
        "extracted_text": "",
        "extracted_fields": {
            "sheets": list(sheets.keys()),
            "preview": sheets,
            "row_count": row_count,
        },
        "page_count": None,
    }


def _parse_xml(content: bytes) -> dict[str, Any]:
    from lxml import etree

    root = etree.fromstring(content)
    tag = etree.QName(root).localname if root.tag.startswith("{") else root.tag
    ns = root.nsmap
    children = [
        etree.QName(c).localname if c.tag.startswith("{") else c.tag
        for c in root
    ]

    return {
        "extracted_text": etree.tostring(root, encoding="unicode", method="text")[:10000],
        "extracted_fields": {
            "root_tag": tag,
            "namespaces": {k: v for k, v in ns.items() if k},
            "child_tags": list(dict.fromkeys(children)),  # уникальные, порядок сохранён
        },
        "page_count": None,
    }


def _parse_docx(content: bytes) -> dict[str, Any]:
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_count = len(doc.tables)

    return {
        "extracted_text": "\n".join(paragraphs),
        "extracted_fields": {"tables_count": tables_count},
        "page_count": None,
    }


def _parse_csv(content: bytes) -> dict[str, Any]:
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    headers = rows[0] if rows else []

    return {
        "extracted_text": text[:10000],
        "extracted_fields": {
            "headers": headers,
            "row_count": len(rows),
        },
        "page_count": None,
    }


def _parse_text(content: bytes) -> dict[str, Any]:
    text = content.decode("utf-8", errors="replace")
    return {
        "extracted_text": text[:50000],
        "extracted_fields": {},
        "page_count": None,
    }
